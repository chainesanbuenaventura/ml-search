import os
import sys
import re
import pandas as pd
import glob
import numpy as np
import matplotlib.pyplot as plt

from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
from tkinter import Tcl
from urllib.parse import urlparse
from pathlib import Path
from fuzzywuzzy import fuzz

sys.path.append("../DataFunctions")
sys.path.append("pipeline/Functions/DataFunctions")
import ElasticFunctions as ef
import MLFlowFunctions as mf


class PCRParagraphExtractor: 
    def __init__(self, file_path):        
        self.file_path = file_path
        self.document = Document(file_path)
        self.curr_header = ''
        self.curr_subheader = ''
        self.appendix_found = False
        self.project_found = False
        self.assessment_found = False
        
        self.paragraphs = []
        self.main_headers = []
        self.sub_headers = []
        
    def is_header(self, paragraph_obj):
        try:
            if (paragraph_obj.style.font.bold 
                and paragraph_obj.style.font.size.pt == 11 
                and paragraph_obj.text.isupper()
                and "BASIC DATA" not in paragraph_obj.text):
                return True 
        except:
            return False 
        return False 
    
    def is_subheader(self, paragraph_obj):
        try:
            if ((paragraph_obj.style.name == 'Title'
                or paragraph_obj.style.name == 'Heading 1'
                or paragraph_obj.style.name == 'List Paragraph')
                and not paragraph_obj.text.isupper()
                and len(paragraph_obj.text.split(' ')) < 10):
                    return True
        except:
            return False 
        return False
    
    def is_paragraph_ok(self, paragraph_obj):
        try:
            #paragraph should not be empty
            if len(paragraph_obj.text.strip()) <= 10:
                return False
            #paragrah should not be a footer
            elif re.search("^\d+\s(.+)", paragraph_obj.text): 
                return False
            #paragraph should not have a normal style
            elif paragraph_obj.style.name == 'Normal':
                return False
            #paragraph should have a font size of 11 
            elif paragraph_obj.style.font.size.pt < 11:
                return False 
        except:
            #paragraph should have more than 10 words 
            if len(paragraph_obj.text.split(' ')) < 10:
                return False
        return True 
    
    def is_in_appendix_section(self, paragraph_obj):
        if (("ISSUES, LESSONS" in self.curr_header
             or "OVERALL ASSESSMENT" in self.curr_header) 
            and "Appendix 1" in paragraph_obj.text):
            return True
        return False
    
    def is_in_project_section(self):
        project_headers = ["PROJECT DESCRIPTION", "PROGRAM DESCRIPTION"]
        for project_header in project_headers:
            if (fuzz.WRatio(project_header, self.curr_header) > 90):
                return True
        return False
    
    def is_in_assessment_section(self):
        assessment_headers = ["ISSUES, LESSONS AND RECOMMENDATIONS", 
                              "OVERALL ASSESSMENT AND RECOMMENDATIONS"]
        for assessment_header in assessment_headers:
            if (fuzz.WRatio(assessment_header, self.curr_header) > 90):
                return True
        return False
        
    def is_main_paragraph(self, paragraph_obj):
        if (paragraph_obj.style.name == 'List Paragraph'
            or (re.search("^\d+\.\s(.+)",paragraph_obj.text)
                and len(paragraph_obj.text.split(' ')) > 50)
            or len(self.paragraphs) == 0):
                return True
        return False
    
    def extract_paragraphs(self):
        print(f"Extracting {self.file_path}")
        for paragraph in self.document.paragraphs: 
            
            if self.is_header(paragraph):
                if self.project_found and not self.appendix_found:
                    print(f"\t In {paragraph.text }")
                self.curr_header = paragraph.text 
                self.curr_subheader = ''
                continue
            elif self.is_subheader(paragraph):
                self.curr_subheader = paragraph.text 
                continue 
            
            #Keep on skipping the table of contents, basic data page until we are in the project section
            if not self.project_found:
                if self.is_in_project_section(): 
                    print(f"\t In {self.curr_header}")
                    self.project_found = True
                else:
                    continue 
            
            if not self.assessment_found:
                if self.is_in_assessment_section():
                    self.assessment_found = True
            
            if (self.is_in_appendix_section(paragraph)
                or (self.assessment_found 
                    and not self.is_in_assessment_section())):
                self.appendix_found = True
            elif not self.is_paragraph_ok(paragraph):
                continue 
                
            if not self.appendix_found and self.curr_header:
                if self.is_main_paragraph(paragraph):
                    self.paragraphs.append(paragraph.text)
                    self.main_headers.append(self.curr_header)
                    self.sub_headers.append(self.curr_subheader)
                else:
                    self.paragraphs[-1] = self.paragraphs[-1] + " " + paragraph.text
                    
        return (self.paragraphs, self.main_headers, self.sub_headers)

# def main():
#     docx_path = "../data/raw_docx/"

#     pcr_data = pd.read_excel("./data/pcr_data_3.xlsx", index_col=0)
#     pcr_data['year'] = [month_year[1] for month_year in pcr_data['Month Year'].str.split(' ')]
#     pcr_data['year'].astype(int)
#     pcr_data.head()

#     project_numbers = (set(pcr_data['Project Number'].values) 
#                        - set(pcr_data.groupby('Project Number').count().query('Titles > 1').index.values))
#     project_number = list(project_numbers)[0]
#     docx_name = glob.glob(docx_path + f"*{project_number}*")
#     pcr_extractor = PCRParagraphExtractor(docx_name[0])
#     paragraphs, main_headers, sub_headers = pcr_extractor.extract_paragraphs()

#     extracted_par_data = []
#     project_numbers = (set(pcr_data['Project Number'].values) 
#                        - set(pcr_data.groupby('Project Number').count().query('Titles > 1').index.values))
#     for project_number in project_numbers:
#         docx_name = glob.glob(docx_path + f"*{project_number}*")
#         try:
#             pcr_extractor = PCRParagraphExtractor(docx_name[0])
#             paragraphs, main_headers, sub_headers = pcr_extractor.extract_paragraphs()

#             extracted_par_data.append(np.array([[proj_number] * len(paragraphs),
#                                                 paragraphs, 
#                                                 main_headers, 
#                                                 sub_headers]))
#         except:
#             print(f"No Document found for {project_number}")

#     return extracted_par_data
        
